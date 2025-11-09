import pandas as pd
import matplotlib.pyplot as mpl
import time




def Processing_The_Data(State_Names):
    
    print(State_Names.head(10))
    State_Names.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Data.xlsx", index=False)
    Max_tem_mean_monthwise = State_Names.groupby(["Year", "Month"])["Maximum_Temperature"].mean().reset_index()
    Min_tem_mean_monthwise = State_Names.groupby(["Year", "Month"])["Minimum_Temperature"].mean().reset_index()
    Total_rain_fall_monthwise = State_Names.groupby(["Year", "Month"])["Rainfall"].sum().reset_index()
    Level_Monthly = State_Names.groupby(["Year", "Month"])["Level"].mean().reset_index()
    Yield_Monthly = State_Names.groupby(["Year", "Month"])["Yield"].mean().reset_index()
    





    Max_tem_mean_yearwise = Max_tem_mean_monthwise.groupby("Year")["Maximum_Temperature"].mean().reset_index()
    Monthly = pd.merge(Max_tem_mean_monthwise,Min_tem_mean_monthwise)
    Monthly.rename(columns={"Maximum_Temperature":"Maximum_Montly_Mean_Temperature","Minimum_Temperature":"Minimum_Monthly_Mean_Temperature"},inplace=True)

    Min_tem_mean_yearwise = Min_tem_mean_monthwise.groupby("Year")["Minimum_Temperature"].mean().reset_index()
    Yearly = pd.merge(Max_tem_mean_yearwise,Min_tem_mean_yearwise)
    Yearly.rename(columns={"Maximum_Temperature":"Maximum_Yearly_Mean_Montly_Mean_Temperature","Minimum_Temperature":"Minimum_Yearly_Mean_Monthly_Mean_Temperature"},inplace=True)

    Total_rain_fall_yearly = Total_rain_fall_monthwise.groupby("Year")["Rainfall"].sum().reset_index()

    
    Monthly = pd.merge(Monthly,Total_rain_fall_monthwise)
    Monthly.rename(columns={"Rainfall":"Total_rain_fall_monthwise"},inplace=True)

    Yearly = pd.merge(Yearly,Total_rain_fall_yearly)
    Yearly.rename(columns={"Rainfall":"Total_rain_fall_yearly"},inplace=True)

    Level_Yearly = Level_Monthly.groupby("Year")["Level"].mean().reset_index()

    print("<----------------=================================---------------->")
    Monthly = pd.merge(Monthly,Level_Monthly)
    Monthly.rename(columns={"Level":"Monthly_Level"},inplace=True)

    Yearly = pd.merge(Yearly,Level_Yearly)
    Yearly.rename(columns={"Level":"Yearly_Level"},inplace=True)

    Yield_Yearly = Yield_Monthly.groupby("Year")["Yield"].mean().reset_index()

    Monthly = pd.merge(Monthly,Yield_Monthly)
    Monthly.rename(columns={"Yield":"Monthly_Yield"},inplace=True)
    print(Monthly)
    Monthly.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Monthly_Data.xlsx", index=False)


    Yearly = pd.merge(Yearly,Yield_Yearly)
    Yearly.rename(columns={"Yield":"Yearly_Yield"},inplace=True)
    print(Yearly)
    Yearly.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Yearly_Data.xlsx", index=False)

    Month_and_Year= pd.DataFrame()
    Month_and_Year["Month-Year"] = Monthly["Month"].astype(str) +"-"+Monthly["Year"].astype(str)





    
    mpl.plot(Yearly["Year"], Yearly["Total_rain_fall_yearly"], marker='o')
    mpl.xticks(Yearly["Year"], rotation=45, ha='right')
    mpl.title(f"Total Rainfall of {state_name}")
    mpl.xlabel("<-------Year------->")
    mpl.ylabel("<-------Rainfall------->")
    mpl.legend()
    mpl.grid(True)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Total_Rainfall.png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(2)           
    mpl.close()



    mpl.plot(Yearly["Year"], Yearly["Maximum_Yearly_Mean_Montly_Mean_Temperature"], marker='o',color = "r")
    mpl.xticks(Yearly["Year"], rotation=45, ha='right')
    mpl.title(f"Maximum Yearly Temperature of {state_name}")
    mpl.xlabel("<-------Year------->")
    mpl.ylabel("<-------Maximum Temperature------->")
    mpl.legend()
    mpl.grid(True)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Maximum_Temperature_(yearly).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(2)           
    mpl.close()



    mpl.plot(Yearly["Year"], Yearly["Minimum_Yearly_Mean_Monthly_Mean_Temperature"], marker='o',color = "orange")
    mpl.xticks(Yearly["Year"], rotation=45, ha='right')
    mpl.title(f"Minimum Yearly Temperature of {state_name}")
    mpl.xlabel("<-------Year------->")
    mpl.ylabel("<-------Minimum Temperature------->")
    mpl.legend()
    mpl.grid(True)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Minimum_Temperature_(yearly).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(2)           
    mpl.close()
   


    
    mpl.plot(Yearly["Year"],Yearly["Yearly_Yield"],marker ="o",color = "g",linewidth=1)
    mpl.xticks(Yearly["Year"], rotation=45, ha='right')
    mpl.title(f"Monthly Average Yield of {state_name}")
    mpl.xlabel("<-------Year-------->")
    mpl.ylabel("<-------Yield------->")
    mpl.grid(True)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Monthly_Average_Yield.png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(2)           
    mpl.close() 



    years_present = sorted(State_Names["Year"].unique())
    expected_years = list(range(2000, 2023))

    # Find missing years
    missing_years = sorted(set(expected_years) - set(years_present))

    if missing_years:
        # Check if all missing years are within the first 3 boundary years (2000–2002)
        allowed_missing = {2000, 2001, 2002,2003}
    
        if all(y in allowed_missing for y in missing_years):
            print(f"ℹ️ Missing only early boundary years {missing_years}. Continuing analysis...")
        else:
            print(f"⚠️ Missing middle or important years {missing_years}. Skipping this state.")
            return
    else:
        print("✅ All years present and continuous. Proceeding...")







    # Classical Decompose
    from statsmodels.tsa.seasonal import seasonal_decompose

    # ✅ Ensure proper datetime index before seasonal decomposition
    Monthly["Date"] = pd.to_datetime(dict(year=Monthly["Year"], month=Monthly["Month"], day=1))
    Monthly = Monthly.set_index("Date")
    Monthly = Monthly.asfreq("MS")  # Monthly Start frequency

    # ✅ Remove missing values in Monthly_Yield
    Monthly = Monthly.dropna(subset=["Monthly_Yield"])
    monthly_yield_clean = Monthly["Monthly_Yield"].dropna()

    # ✅ Handle short or incomplete data safely
    if len(monthly_yield_clean) >= 24:
        decompose_additive = seasonal_decompose(monthly_yield_clean, model="additive", period=12)
        Trend_decompose = decompose_additive.trend
        Seasonal_decompose = decompose_additive.seasonal
        Residual_decompose = decompose_additive.resid
    else:
        print(f"⚠️ Skipping seasonal decomposition for {state_name} (only {len(monthly_yield_clean)} months — need ≥ 24).")
        Trend_decompose = Seasonal_decompose = Residual_decompose = None

    # ✅ Skip plotting if decomposition data missing
    if Trend_decompose is None:
        print(f"⚠️ Not enough data to plot decomposition for {state_name}. Skipping plots.")
        return



    
    mpl.plot(Month_and_Year["Month-Year"],Monthly["Monthly_Yield"], label ="Original",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"Classical Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Classical_Decompose_of_{state_name}(Original).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    mpl.plot(Month_and_Year["Month-Year"],Trend_decompose, label ="Trend",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"Classical Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Classical_Decompose_of_{state_name}(Trend).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    mpl.plot(Month_and_Year["Month-Year"],Seasonal_decompose, label ="Seasonal",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"Classical Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Classical_Decompose_of_{state_name}(Seasonal).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    mpl.plot(Month_and_Year["Month-Year"],Residual_decompose, label ="Residual",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"Classical Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Classical_Decompose_of_{state_name}(Residual).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    #STL Decomposition
    from statsmodels.tsa.seasonal import STL
    STL_Decomposition_additive = STL(Monthly["Monthly_Yield"],period=12)
    result = STL_Decomposition_additive.fit()

    
    mpl.plot(Month_and_Year["Month-Year"],result.observed, label ="Original",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"STL Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\STL_Decompose_of_{state_name}(Original).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    mpl.plot(Month_and_Year["Month-Year"],result.trend, label ="Trend",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"STL Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\STL_Decompose_of_{state_name}(Trend).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    mpl.plot(Month_and_Year["Month-Year"],result.seasonal, label ="Seasonal",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"STL Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\STL_Decompose_of_{state_name}(Seasonal).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    mpl.plot(Month_and_Year["Month-Year"],result.resid, label ="Residual",color ="r")
    mpl.legend(loc ="upper left")
    mpl.title(f"STL Decompose of {state_name}")
    mpl.xlabel("<-------Month-Year------->")
    mpl.ylabel("<-------Monthly_Yield------->")
    mpl.grid(True)
    mpl.xticks(Month_and_Year["Month-Year"][::12],rotation=45)
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\STL_Decompose_of_{state_name}(Residual).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    #stationarity Test:-

    #To perform the Augmented Dicky-fuler test:-

    #ADF test:-
    from statsmodels.tsa.stattools import adfuller

    def call_ADF(source):
        adf_test = adfuller(source.dropna())
        adf_stat = adf_test[0]
        p_val = adf_test[1]
        crit_5 = adf_test[4]['5%']
    
        print("ADF Test Results (5% level):")
        print(f"ADF Statistic: {round(adf_stat, 3)}")
        print(f"p-value: {round(p_val, 3)}")
        print(f"5% Critical Value: {round(crit_5, 3)}")
    
        if p_val < 0.05 and adf_stat < crit_5:
            print("✅ Data is **Stationary** at 5% significance level.(ADF Test)")
        else:
            print("❌ Data is **Non-Stationary** at 5% significance level.(ADF Test)")
            return False
        


    #KPSS Test:-
    from statsmodels.tsa.stattools import kpss

    def call_KPSS(source):
        kpss_test = kpss(source.dropna(), regression="ct")  # 'ct' → trend stationarity
        kpss_stat = kpss_test[0]
        p_val = kpss_test[1]
        crit_5 = kpss_test[3]['5%']

        print("KPSS Test Results (5% level):")
        print(f"KPSS Statistic: {round(kpss_stat, 3)}")
        print(f"p-value: {round(p_val, 3)}")
        print(f"5% Critical Value: {round(crit_5, 3)}")

        if p_val > 0.05 and kpss_stat < crit_5:
            print("✅ Data is **Stationary** at 5% significance level.(KPSS Test)")
        else:
            print("❌ Data is **Non-Stationary** at 5% significance level.(KPSS Test)")
            return False
        

    

    
    ADF_Return = call_ADF(Monthly["Monthly_Yield"])
    KPSS_Return =call_KPSS(Monthly["Monthly_Yield"])
    if ADF_Return == False or KPSS_Return == False:
        Monthly["Yield_Diff"] = Monthly["Monthly_Yield"].diff()
        Monthly["Yield_Diff"] = Monthly["Yield_Diff"].dropna()
        call_ADF(Monthly["Yield_Diff"])
        call_KPSS(Monthly["Yield_Diff"])
    else:
        Monthly["Yield_Diff"] =Monthly["Monthly_Yield"]
   


    # -----------------------------
    # Step 1: Detect outliers using IQR
    # -----------------------------
    Q1 = Monthly["Yield_Diff"].quantile(0.25)
    Q3 = Monthly["Yield_Diff"].quantile(0.75)
    IQR = Q3 - Q1

    # Define bounds
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR

    # Find outliers
    outliers = Monthly[(Monthly["Yield_Diff"] < lower_bound) | (Monthly["Yield_Diff"] > upper_bound)]
    print(f"Number of outliers detected: {len(outliers)}")
    print("Outlier rows:\n", outliers)
    outliers.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_outliers.xlsx", index=False)

    # -----------------------------
    # Step 2: Remove outliers
    # -----------------------------
    Monthly_clean = Monthly[~((Monthly["Yield_Diff"] < lower_bound) | (Monthly["Yield_Diff"] > upper_bound))].copy()
    print(f"Rows after outlier removal: {Monthly_clean.shape[0]}")

    # -----------------------------
    # Step 3: Reset index
    # -----------------------------
    Monthly_clean = Monthly_clean.reset_index(drop=True)

    print(Monthly_clean)
    Monthly_clean.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Monthly_clean_Data.xlsx", index=False)
   

















    # Split train and test by year
    max_year = Monthly_clean["Year"].max()
    test_years = [max_year - 1, max_year]

    train = Monthly_clean[~Monthly_clean["Year"].isin(test_years)]
    test = Monthly_clean[Monthly_clean["Year"].isin(test_years)]

    print(f"📅 Training years: {sorted(train['Year'].unique())}")
    print(f"🧪 Testing years: {sorted(test['Year'].unique())}")

    y_train = train["Monthly_Yield"]
    y_test = test["Monthly_Yield"]

    exog_vars = ["Total_rain_fall_monthwise", "Monthly_Level", "Maximum_Montly_Mean_Temperature","Minimum_Monthly_Mean_Temperature"]
    X_train = train[exog_vars]
    X_test = test[exog_vars]
    
    

    from statsmodels.tsa.statespace.sarimax import SARIMAX

    model = SARIMAX(y_train, exog=X_train, order=(1, 1, 1), seasonal_order=(1,1,1,12))
    results = model.fit(disp=False)
    print(results.summary())


    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    summary_text = results.summary().as_text()

    doc = Document()
    doc.add_heading(f"SARIMAX Model Summary - {state_name}", level=1)

    # Create 1-cell table to hold the summary text
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p = cell.add_paragraph(summary_text)
    run = p.runs[0]
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ✅ Add black border to the cell
    tcPr = cell._element.get_or_add_tcPr()
    for border in ["top", "left", "bottom", "right"]:
        border_el = OxmlElement(f"w:{border}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "4")
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), "000000")
        tcPr.append(border_el)

    doc.save(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_Summary.docx")
    
    forecast = results.predict(start=len(y_train), end=len(y_train) + len(y_test) - 1, exog=X_test)
    forecast = forecast.reset_index(drop=True)


    from sklearn.metrics import mean_absolute_error, mean_squared_error

    mae = mean_absolute_error(y_test, forecast)
    mse = mean_squared_error(y_test, forecast)
    print(f"MAE: {mae:.4f}")
    print(f"MSE: {mse:.4f}")
    




    # Make sure forecast index matches the test index (for 2021–2022)
    forecast = results.predict(start=y_test.index[0], end=y_test.index[-1], exog=X_test)

    # Combine original and forecast data into one DataFrame
    compare_df = pd.DataFrame({
        "Date": y_test.index,
        "Actual_Monthly_Yield": y_test.values,
        "Forecasted_Monthly_Yield": forecast.values
    })

    # Reset index for pretty printing
    compare_df = compare_df.reset_index(drop=True)

    # Print full comparison
    print("\n📊 Comparison: Actual vs Forecasted Monthly Yield (2021–2022)\n")
    print(compare_df)
    compare_df.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_compare.xlsx", index=False)



    import statsmodels.api as sm

    # Assuming `results` is your SARIMAX fitted model
    residuals = results.resid
    
    # 1. Plot residuals
    mpl.figure(figsize=(12,4))
    mpl.plot(residuals)
    mpl.title(f"Residuals over Time ({state_name})")
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Residuals_over_Time_({state_name}).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    # 2. Histogram + density of residuals
    mpl.figure(figsize=(8,4))
    residuals.plot(kind='hist', bins=30, density=True)
    residuals.plot(kind='kde')
    mpl.title(f"Residual Distribution ({state_name})")
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Residuals_Distribution({state_name}).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    # 3. ACF/PACF of residuals
    fig, ax = mpl.subplots(2,1, figsize=(12,6))
    sm.graphics.tsa.plot_acf(residuals, ax=ax[0])
    sm.graphics.tsa.plot_pacf(residuals, ax=ax[1])
    mpl.title(f"ACF and PACF of residual ({state_name})")
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\ACF_and_PACF_of_residual({state_name}).png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    import numpy as np
    mape = np.mean(np.abs((y_test - forecast)/y_test)) * 100
    print("MAPE:", mape, "%")

    results = model.fit() 
    from statsmodels.tsa.statespace.sarimax import SARIMAX





    # ✅ Create proper datetime column
    Monthly["Date"] = pd.to_datetime(dict(year=Monthly["Year"], month=Monthly["Month"], day=1))
    Monthly.set_index("Date", inplace=True)


    # Drop rows with invalid/missing dates
    monthly_clean = Monthly[Monthly.index.notna()]


    # ✅ Automatically detect available range
    start_year = monthly_clean.index.year.min()
    end_year = monthly_clean.index.year.max()

    print(f"Available data range: {Monthly.index.min().year}–{Monthly.index.max().year}")


    # ✅ Use last 5 years if possible
    if end_year - start_year >= 5:
        recent_years = monthly_clean.loc[f"{end_year-4}-01-01":f"{end_year}-12-31"]
    else:
        recent_years = monthly_clean.copy()

    print(f"Recent subset used for exog pattern: {recent_years.index.min().year}–{recent_years.index.max().year}")

    # ✅ Group by month and compute averages
    if not recent_years.empty:
        monthly_exog_pattern = recent_years.groupby(recent_years.index.month)[
            ["Total_rain_fall_monthwise", "Monthly_Level", 
            "Maximum_Montly_Mean_Temperature", "Minimum_Monthly_Mean_Temperature"]
        ].mean()
    else:
        raise ValueError("No data available in recent_years subset!")

    # ✅ Fill missing months (ensure 1–12 present)
    monthly_exog_pattern = monthly_exog_pattern.reindex(range(1, 13)).interpolate()

    # ✅ Create 2023 exog dataframe with monthly-varying values
    forecast_index_2023 = pd.date_range(start="2023-01-01", periods=12, freq="MS")

    # Ensure exog pattern cycles correctly month-wise (1→Jan, 2→Feb, etc.)
    exog_2023 = pd.DataFrame({
        "Total_rain_fall_monthwise": monthly_exog_pattern.loc[forecast_index_2023.month, "Total_rain_fall_monthwise"].values,
        "Monthly_Level": monthly_exog_pattern.loc[forecast_index_2023.month, "Monthly_Level"].values,
        "Maximum_Montly_Mean_Temperature": monthly_exog_pattern.loc[forecast_index_2023.month, "Maximum_Montly_Mean_Temperature"].values,
        "Minimum_Monthly_Mean_Temperature": monthly_exog_pattern.loc[forecast_index_2023.month, "Minimum_Monthly_Mean_Temperature"].values
    }, index=forecast_index_2023)

    # Format index to show only YYYY-MM when printing
    exog_2023.index = exog_2023.index.strftime("%Y-%m")

    print("\n✅ Generated realistic 2023 exogenous forecast (month-wise):")
    print(exog_2023)
    exog_2023.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_exog_2023.xlsx", index=True)

    # ✅ Extend exogenous data pattern for 2024 (copied or slightly varied from 2023)
    exog_2024 = exog_2023.copy()

    # Option 1: Keep same monthly pattern as 2023
    exog_2024.index = pd.date_range("2024-01", "2024-12", freq="MS")

    # Option 2 (optional): add small random variation for realism
    import numpy as np
    exog_2024["Total_rain_fall_monthwise"] *= np.random.uniform(0.9, 1.1, size=12)
    exog_2024["Monthly_Level"] *= np.random.uniform(0.95, 1.05, size=12)
    exog_2024["Maximum_Montly_Mean_Temperature"] *= np.random.uniform(0.98, 1.02, size=12)
    exog_2024["Minimum_Monthly_Mean_Temperature"] *= np.random.uniform(0.98, 1.02, size=12)

    print("\n✅ Generated realistic 2024 exogenous forecast:")
    print(exog_2024)
    exog_2024.to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_exog_2024.xlsx", index=True)

    # Combine both years
    exog_2023_2024 = pd.concat([exog_2023, exog_2024])




    # Forecast 12 months ahead (Jan–Dec 2023)
    forecast_2023 = results.get_forecast(steps=12, exog=exog_2023)
    pred_mean = forecast_2023.predicted_mean
    conf_int = forecast_2023.conf_int()

    # ✅ Fix the index to reflect 2023 dates
    pred_mean.index = pd.date_range(start="2023-01", periods=12, freq="MS")
    conf_int.index = pred_mean.index



    # Extend exogenous variables for 2024 (same pattern as 2023 for simplicity)
    exog_2024 = exog_2023.copy()
    exog_2024.index = pd.date_range("2024-01", "2024-12", freq="MS")

    # Combine both
    exog_2023_2024 = pd.concat([exog_2023, exog_2024])

    # Generate forecast for 2023–2024
    forecast_2023_2024 = results.get_forecast(steps=24, exog=exog_2023_2024)
    predicted_values = forecast_2023_2024.predicted_mean

    # ✅ Assign proper date index
    forecast_index = pd.date_range("2023-01", periods=24, freq="MS")
    predicted_values.index = forecast_index

    # ✅ Split 2023 and 2024
    forecast_2023 = predicted_values[:12]
    forecast_2024 = predicted_values[12:]

    # ✅ Convert to DataFrame with Month & Year
    forecast_2023_df = forecast_2023.reset_index()
    forecast_2023_df.columns = ["Date", "Forecasted_Monthly_Yield"]
    forecast_2023_df["Year"] = forecast_2023_df["Date"].dt.year
    forecast_2023_df["Month"] = forecast_2023_df["Date"].dt.strftime("%b")

    forecast_2024_df = forecast_2024.reset_index()
    forecast_2024_df.columns = ["Date", "Forecasted_Monthly_Yield"]
    forecast_2024_df["Year"] = forecast_2024_df["Date"].dt.year
    forecast_2024_df["Month"] = forecast_2024_df["Date"].dt.strftime("%b")

    # ✅ Nicely formatted output
    print("\n📈 Forecasted Monthly Yield for 2023:")
    print(forecast_2023_df[["Year", "Month", "Forecasted_Monthly_Yield"]].to_string(index=False))
    forecast_2023_df[["Year", "Month", "Forecasted_Monthly_Yield"]].to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_forecast_2023.xlsx", index=False)

    print("\n📈 Forecasted Monthly Yield for 2024:")
    print(forecast_2024_df[["Year", "Month", "Forecasted_Monthly_Yield"]].to_string(index=False))
    forecast_2024_df[["Year", "Month", "Forecasted_Monthly_Yield"]].to_excel(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\{state_name}_forecast_2024.xlsx", index=False)

    # 🧭 Fix missing datetime index for plotting
    train["Date"] = pd.to_datetime(dict(year=train["Year"], month=train["Month"], day=1))
    test["Date"] = pd.to_datetime(dict(year=test["Year"], month=test["Month"], day=1))

    y_train.index = train["Date"]
    y_test.index = test["Date"]

    """

    # =====================================================
    # 📊 PLOT: Train vs Test vs Forecast (2023–2024)
    # =====================================================
    import matplotlib.dates as mdates
    
    if len(y_train) < 5 or len(y_test) == 0:
        print(f"⚠️ Not enough data for {state}, skipped.")


    mpl.figure(figsize=(12, 6))

    # ✅ Plot training data
    mpl.plot(y_train.index, y_train, label="Training Data", color="blue", linewidth=2)

    # ✅ Plot actual testing data
    mpl.plot(y_test.index, y_test, label="Actual Test Data", color="orange", linewidth=2)

    # ✅ Plot predicted (forecasted) values for test period
    mpl.plot(y_test.index, forecast[:len(y_test)], label="Predicted Test (Model Output)", color="red", linestyle="--", linewidth=2)

    # ✅ Plot forecasted 2023–2024 values
    mpl.plot(predicted_values.index, predicted_values, label="Forecast (2023–2024)", color="green", linewidth=2)

    # ✅ Add confidence interval shading for forecast period
    conf_int = forecast_2023_2024.conf_int(alpha=0.05)
    conf_int.index = predicted_values.index
    mpl.fill_between(predicted_values.index,
                 conf_int.iloc[:, 0],
                 conf_int.iloc[:, 1],
                 color="lightgreen", alpha=0.3)

    # ✅ Formatting
    mpl.title(f"Training vs Testing vs Forecasted Monthly Yield of {state_name}")
    mpl.xlabel("<-------Year------->")
    mpl.ylabel("<-------Yield------->")
    mpl.gca().xaxis.set_major_locator(mdates.YearLocator())
    mpl.gca().xaxis.set_major_formatter(mdates.DateFormatter("%Y"))
    mpl.xticks(rotation=45)
    mpl.grid(True, linestyle="--", alpha=0.5)
    mpl.legend()
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Training_vs_Testing_vs_Forecasted_Monthly_Yield_of_{state_name}.png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()
    """







    # =====================================================
    # 📊 Separate Plots for 2023 and 2024 with Confidence Intervals
    # =====================================================

    
    import matplotlib.dates as mdates

    # ✅ Extract confidence intervals and set date index
    conf_int = forecast_2023_2024.conf_int(alpha=0.05)
    conf_int.index = predicted_values.index

    # ✅ Split by year
    pred_2023 = predicted_values["2023"]
    pred_2024 = predicted_values["2024"]
    conf_2023 = conf_int.loc["2023"]
    conf_2024 = conf_int.loc["2024"]

    # ---------- 📈 Plot for 2023 ----------
    mpl.figure(figsize=(10, 5))
    mpl.plot(pred_2023.index, pred_2023.values, label="Forecast 2023", color="royalblue", linewidth=2)
    mpl.fill_between(pred_2023.index, conf_2023.iloc[:, 0], conf_2023.iloc[:, 1],
                 color="lightblue", alpha=0.3)
    mpl.title(f"Forecasted Monthly Yield for 2023 (95% Confidence Interval)({state_name})")
    mpl.xlabel("<-------Month------->")
    mpl.ylabel("<-------Predicted Yield------->")
    mpl.gca().xaxis.set_major_formatter(mdates.DateFormatter("%b"))
    mpl.grid(True, linestyle="--", alpha=0.5)
    mpl.legend()
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Forecasted_Monthly_Yield_for_2023_(95%_Confidence_Interval)_{state_name}.png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()

    # ---------- 📈 Plot for 2024 ----------
    mpl.figure(figsize=(10, 5))
    mpl.plot(pred_2024.index, pred_2024.values, label="Forecast 2024", color="green", linewidth=2)
    mpl.fill_between(pred_2024.index, conf_2024.iloc[:, 0], conf_2024.iloc[:, 1],
                 color="lightgreen", alpha=0.3)
    mpl.title(f"Forecasted Monthly Yield for 2024 (95% Confidence Interval)({state_name})")
    mpl.xlabel("<-------Month------->")
    mpl.ylabel("<-------Predicted Yield------->")
    mpl.gca().xaxis.set_major_formatter(mdates.DateFormatter("%b"))
    mpl.grid(True, linestyle="--", alpha=0.5)
    mpl.legend()
    mpl.tight_layout()
    mpl.savefig(f"C:\\Users\\PIPLI KHANRA\\Desktop\\ISI_Internship\\Forecasted_Monthly_Yield_for_2024_(95%_Confidence_Interval)_{state_name}.png",dpi = 300)
    mpl.show(block=False)
    mpl.pause(0.1)
    time.sleep(10)           
    mpl.close()










Main_Data = pd.read_csv("merged_rabi_rice_reservoir.csv")
State_Names =Main_Data['state_name'].unique()
print(f"State name:-\n{State_Names}")

State_Data = {}  # dictionary to store all state data

for state in State_Names:
    # Create dataframe for each state
    raw_data = Main_Data[Main_Data["state_name"] == state].copy()

    df = pd.DataFrame()
    df["Crop_name"] =raw_data["crop_name"]
    df["Year-Month-Date"] =pd.to_datetime(raw_data["temperature_recorded_date"]) 
    df["Year"] = df["Year-Month-Date"].dt.year
    df["Month"] = df["Year-Month-Date"].dt.month
    df["Date"] = df["Year-Month-Date"].dt.date
    df["Maximum_Temperature"] = raw_data["state_temperature_max_val"]
    df["Minimum_Temperature"] = raw_data["state_temperature_min_val"]
    df["Maximum_Temperature"] =raw_data["state_temperature_max_val"]
    df["Minimum_Temperature"] =raw_data["state_temperature_min_val"]
    df["Rainfall"] =raw_data["state_rainfall_val"]
    df["Yield"] =raw_data["yield"]
    df["FRL"] =raw_data["FRL"]
    df["Live_Cap_FRL"] =raw_data["Live Cap FRL"]
    df["Level"] =raw_data["Level"]


# Store each state's DataFrame in dictionary (with underscores in name)
    key_name = state.replace(" ", "_")
    State_Data[key_name] = df

# --- Automatically loop over all states and run your function ---
for state_name, state_df in State_Data.items():
    print(f"\n==============================")
    print(f"🚀 Current State : {state_name}")
    print(f"==============================")
    Processing_The_Data(state_df)
   
   
